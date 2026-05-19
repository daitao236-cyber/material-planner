# services/document_builder.py
# 文档构建服务

import logging
from typing import Optional, List, Dict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from domain.models import MaterialPlan, PlanSection, PlanItem, FemaleStyle, TimelineItem

logger = logging.getLogger(__name__)


class DocumentBuilder:
    """文档构建服务"""

    def __init__(self):
        self._styles_loaded = False

    def build_markdown(self, plan: MaterialPlan) -> str:
        """
        构建Markdown格式文档

        Args:
            plan: 素材规划对象

        Returns:
            str: Markdown格式内容
        """
        lines = []

        # 标题
        lines.append(f"# {plan.title}\n")
        lines.append(f"**赛季**: {plan.season}\n")
        lines.append(f"**生成时间**: {plan.created_at.strftime('%Y-%m-%d %H:%M')}\n")
        lines.append(f"**AI素材占比**: {plan.ai_ratio * 100:.0f}%\n")
        lines.append(f"**常规素材占比**: {plan.normal_ratio * 100:.0f}%\n")

        # 核心主旨
        if plan.core_themes:
            lines.append("\n## 一、核心主旨\n")
            for i, theme in enumerate(plan.core_themes, 1):
                lines.append(f"{i}. {theme}")

        # 板块分类
        if plan.sections:
            lines.append("\n## 二、板块分类\n")
            for section in plan.sections:
                lines.append(f"\n### {section.name} (优先级: {section.priority})\n")
                if section.description:
                    lines.append(f"{section.description}\n")

                for item in section.items:
                    lines.append(f"- **沟通口径**: {item.沟通口径}")
                    lines.append(f"  - **创意延展**: {item.创意延展}")
                    lines.append(f"  - **优先级**: {item.优先级}")
                    if item.预估效果:
                        lines.append(f"  - **预估效果**: {item.预估效果}")
                    lines.append("")

        # AI漫剧专项
        if plan.male_content or plan.female_content:
            lines.append("\n## 三、AI漫剧专项\n")

            if plan.male_content:
                lines.append("\n### 男频题材方向\n")
                for content in plan.male_content:
                    lines.append(f"- {content}")

            if plan.female_content:
                lines.append("\n### 女频题材方向\n")
                for content in plan.female_content:
                    lines.append(f"- {content}")

        # 女性向素材
        if plan.female_styles:
            lines.append("\n## 四、女性向素材\n")
            for style in plan.female_styles:
                lines.append(f"\n### {style.type}\n")
                lines.append(f"**目标人群**: {style.target_audience}\n")
                lines.append("**内容方向**:\n")
                for direction in style.content_directions:
                    lines.append(f"- {direction}")
                if style.recommended_elements:
                    lines.append("**推荐元素**:\n")
                    for elem in style.recommended_elements:
                        lines.append(f"- {elem}")

        # 制作周期
        if plan.timeline:
            lines.append("\n## 五、制作周期\n")
            for item in plan.timeline:
                lines.append(f"\n### {item.phase}\n")
                lines.append(f"**时间**: {item.start_date} ~ {item.end_date}\n")
                lines.append("**执行动作**:\n")
                for action in item.actions:
                    lines.append(f"- {action}")
                lines.append("**输出结果**:\n")
                for output in item.outputs:
                    lines.append(f"- {output}")

        # 数据洞察
        if plan.data_insights:
            lines.append("\n## 六、数据洞察\n")
            for key, value in plan.data_insights.items():
                lines.append(f"- **{key}**: {value}")

        # 热门趋势
        if plan.trends:
            lines.append("\n## 七、热门趋势\n")
            for trend in plan.trends[:15]:
                platform = trend.get("platform", "")
                keyword = trend.get("keyword", "")
                heat = trend.get("heat", 0)
                lines.append(f"- [{platform}] {keyword} (热度: {heat:.0f})")

        return "\n".join(lines)

    def build_word(self, plan: MaterialPlan) -> bytes:
        """
        构建Word格式文档

        Args:
            plan: 素材规划对象

        Returns:
            bytes: Word文档字节数据
        """
        doc = Document()

        # 设置标题样式
        self._apply_title_style(doc)

        # 标题
        title = doc.add_heading(plan.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        doc.add_paragraph(f"赛季: {plan.season}")
        doc.add_paragraph(f"生成时间: {plan.created_at.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"AI素材占比: {plan.ai_ratio * 100:.0f}%")
        doc.add_paragraph(f"常规素材占比: {plan.normal_ratio * 100:.0f}%")

        # 核心主旨
        if plan.core_themes:
            doc.add_heading("一、核心主旨", 1)
            for i, theme in enumerate(plan.core_themes, 1):
                p = doc.add_paragraph(f"{i}. {theme}")

        # 板块分类
        if plan.sections:
            doc.add_heading("二、板块分类", 1)
            for section in plan.sections:
                doc.add_heading(section.name, 2)
                p = doc.add_paragraph(f"优先级: {section.priority}")
                if section.description:
                    doc.add_paragraph(section.description)

                # 条目表格
                if section.items:
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '沟通口径'
                    hdr_cells[1].text = '创意延展'
                    hdr_cells[2].text = '优先级'
                    hdr_cells[3].text = '预估效果'

                    for item in section.items:
                        row_cells = table.add_row().cells
                        row_cells[0].text = item.沟通口径
                        row_cells[1].text = item.创意延展
                        row_cells[2].text = item.优先级
                        row_cells[3].text = item.预估效果 or "-"

        # AI漫剧专项
        if plan.male_content or plan.female_content:
            doc.add_heading("三、AI漫剧专项", 1)

            if plan.male_content:
                doc.add_heading("男频题材方向", 2)
                for content in plan.male_content:
                    doc.add_paragraph(content, style='List Bullet')

            if plan.female_content:
                doc.add_heading("女频题材方向", 2)
                for content in plan.female_content:
                    doc.add_paragraph(content, style='List Bullet')

        # 女性向素材
        if plan.female_styles:
            doc.add_heading("四、女性向素材", 1)
            for style in plan.female_styles:
                doc.add_heading(style.type, 2)
                doc.add_paragraph(f"目标人群: {style.target_audience}")
                
                p = doc.add_paragraph()
                p.add_run("内容方向:").bold = True
                for direction in style.content_directions:
                    doc.add_paragraph(direction, style='List Bullet')

                if style.platform_suggestions:
                    p = doc.add_paragraph()
                    p.add_run("推荐平台:").bold = True
                    p.add_run(", ".join(style.platform_suggestions))

        # 制作周期
        if plan.timeline:
            doc.add_heading("五、制作周期", 1)
            for item in plan.timeline:
                doc.add_heading(item.phase, 2)
                doc.add_paragraph(f"时间: {item.start_date} ~ {item.end_date}")

                p = doc.add_paragraph()
                p.add_run("执行动作:").bold = True
                for action in item.actions:
                    doc.add_paragraph(action, style='List Bullet')

                p = doc.add_paragraph()
                p.add_run("输出结果:").bold = True
                for output in item.outputs:
                    doc.add_paragraph(output, style='List Bullet')

        # 数据洞察
        if plan.data_insights:
            doc.add_heading("六、数据洞察", 1)
            for key, value in plan.data_insights.items():
                doc.add_paragraph(f"{key}: {value}")

        # 热门趋势
        if plan.trends:
            doc.add_heading("七、热门趋势", 1)
            for trend in plan.trends[:20]:
                platform = trend.get("platform", "")
                keyword = trend.get("keyword", "")
                heat = trend.get("heat", 0)
                doc.add_paragraph(f"[{platform}] {keyword} (热度: {heat:.0f})")

        # 保存到字节流
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _apply_title_style(self, doc: Document):
        """应用标题样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)

    def export(
        self,
        plan: MaterialPlan,
        output_path: str | Path,
        format: str = "md"
    ) -> str:
        """
        导出文档

        Args:
            plan: 素材规划对象
            output_path: 输出路径
            format: 输出格式 (md/docx)

        Returns:
            str: 保存的文件路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if format == "md":
            content = self.build_markdown(plan)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            logger.info(f"Markdown文档已保存: {output_path}")

        elif format == "docx":
            content = self.build_word(plan)
            with open(output_path, "wb") as f:
                f.write(content)
            logger.info(f"Word文档已保存: {output_path}")

        else:
            raise ValueError(f"不支持的格式: {format}")

        return str(output_path)

    def export_multi_format(
        self,
        plan: MaterialPlan,
        output_dir: str | Path,
        formats: Optional[List[str]] = None
    ) -> Dict[str, str]:
        """
        导出多种格式

        Args:
            plan: 素材规划对象
            output_dir: 输出目录
            formats: 格式列表

        Returns:
            Dict[str, str]: 格式到文件路径的映射
        """
        if formats is None:
            formats = ["md", "docx"]

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # 生成文件名
        timestamp = plan.created_at.strftime("%Y%m%d_%H%M%S")
        base_name = f"素材规划_{timestamp}"

        results = {}

        for fmt in formats:
            if fmt == "md":
                file_path = output_dir / f"{base_name}.md"
                self.export(plan, file_path, "md")
                results["markdown"] = str(file_path)

            elif fmt == "docx":
                file_path = output_dir / f"{base_name}.docx"
                self.export(plan, file_path, "docx")
                results["word"] = str(file_path)

        return results


def create_document_builder() -> DocumentBuilder:
    """创建文档构建服务"""
    return DocumentBuilder()